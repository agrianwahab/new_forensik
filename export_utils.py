"""
Export Utilities Module for Forensic Image Analysis System
Contains functions for exporting results to various formats (DOCX, PDF, PNG, TXT)
"""

import os
import io
import subprocess
import platform
import shutil
from datetime import datetime
from PIL import Image
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

# DOCX imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import warnings
warnings.filterwarnings('ignore')

# ======================= Main Export Functions =======================

def export_complete_package(original_pil, analysis_results, base_filename="forensic_analysis"):
    """Export complete analysis package (PNG, PDF visualization, DOCX report, PDF report)"""
    print(f"\n{'='*80}")
    print("📦 CREATING COMPLETE EXPORT PACKAGE")
    print(f"{'='*80}")
    
    export_files = {}
    
    try:
        # 1. Export PNG visualization
        png_file = f"{base_filename}_visualization.png"
        export_files['png_visualization'] = export_visualization_png(original_pil, analysis_results, png_file)
        
        # 2. Export PDF visualization
        pdf_viz_file = f"{base_filename}_visualization.pdf"
        export_files['pdf_visualization'] = export_visualization_pdf(original_pil, analysis_results, pdf_viz_file)
        
        # 3. Export DOCX report
        docx_file = f"{base_filename}_report.docx"
        export_files['docx_report'] = export_to_advanced_docx(original_pil, analysis_results, docx_file)
        
        # 4. Export PDF report
        pdf_report_file = f"{base_filename}_report.pdf"
        pdf_result = export_report_pdf(docx_file, pdf_report_file)
        if pdf_result:
            export_files['pdf_report'] = pdf_result
        
        # 5. Create summary file
        summary_file = f"{base_filename}_summary.txt"
        export_files['summary'] = create_export_summary(analysis_results, export_files, summary_file)
        
        # 6. Export K-means visualization (if available)
        if 'localization_analysis' in analysis_results:
            kmeans_file = f"{base_filename}_kmeans.jpg"
            from visualization import export_kmeans_visualization
            kmeans_result = export_kmeans_visualization(original_pil, analysis_results, kmeans_file)
            if kmeans_result:
                export_files['kmeans_visualization'] = kmeans_result
        
    except Exception as e:
        print(f"❌ Error during export: {e}")
    
    print(f"\n{'='*80}")
    print("📦 EXPORT PACKAGE COMPLETE")
    print(f"{'='*80}")
    print("📁 Generated Files:")
    
    for file_type, filename in export_files.items():
        if filename and os.path.exists(filename):
            file_size = os.path.getsize(filename)
            print(f"  ✅ {file_type}: {filename} ({file_size:,} bytes)")
        else:
            print(f"  ❌ {file_type}: Failed to create")
    
    print(f"{'='*80}\n")
    
    return export_files

# ======================= Visualization Export Functions =======================

def export_visualization_png(original_pil, analysis_results, output_filename="forensic_analysis.png"):
    """Export visualization to PNG format with high quality"""
    print("📊 Creating PNG visualization...")
    
    try:
        from visualization import visualize_results_advanced
        return visualize_results_advanced(original_pil, analysis_results, output_filename)
    except ImportError:
        print("❌ Visualization module not available")
        return None
    except Exception as e:
        print(f"❌ Error creating PNG visualization: {e}")
        return None

def export_visualization_pdf(original_pil, analysis_results, output_filename="forensic_analysis.pdf"):
    """Export visualization to PDF format"""
    print("📊 Creating PDF visualization...")
    
    try:
        with PdfPages(output_filename) as pdf:
            # Page 1: Main Analysis
            fig1 = plt.figure(figsize=(16, 12))
            gs1 = fig1.add_gridspec(3, 4, hspace=0.3, wspace=0.3)
            fig1.suptitle("Forensic Image Analysis - Main Results", fontsize=16, fontweight='bold')
            
            # Import visualization functions
            from visualization import (create_feature_match_visualization, create_block_match_visualization,
                                     create_frequency_visualization, create_texture_visualization,
                                     create_technical_metrics_plot, create_edge_visualization,
                                     create_illumination_visualization, create_statistical_visualization,
                                     create_quality_response_plot, create_advanced_combined_heatmap)
            
            # Row 1: Core Analysis
            ax1 = fig1.add_subplot(gs1[0, 0])
            ax1.imshow(original_pil)
            ax1.set_title("Original Image", fontsize=12)
            ax1.axis('off')
            
            ax2 = fig1.add_subplot(gs1[0, 1])
            ela_display = ax2.imshow(analysis_results['ela_image'], cmap='hot')
            ax2.set_title(f"ELA (μ={analysis_results['ela_mean']:.1f})", fontsize=12)
            ax2.axis('off')
            plt.colorbar(ela_display, ax=ax2, fraction=0.046)
            
            ax3 = fig1.add_subplot(gs1[0, 2])
            create_feature_match_visualization(ax3, original_pil, analysis_results)
            
            ax4 = fig1.add_subplot(gs1[0, 3])
            create_block_match_visualization(ax4, original_pil, analysis_results)
            
            # Row 2: Advanced Analysis
            ax5 = fig1.add_subplot(gs1[1, 0])
            create_frequency_visualization(ax5, analysis_results)
            
            ax6 = fig1.add_subplot(gs1[1, 1])
            create_texture_visualization(ax6, analysis_results)
            
            ax7 = fig1.add_subplot(gs1[1, 2])
            ghost_display = ax7.imshow(analysis_results['jpeg_ghost'], cmap='hot')
            ax7.set_title(f"JPEG Ghost", fontsize=12)
            ax7.axis('off')
            plt.colorbar(ghost_display, ax=ax7, fraction=0.046)
            
            ax8 = fig1.add_subplot(gs1[1, 3])
            create_technical_metrics_plot(ax8, analysis_results)
            
            # Row 3: Summary
            ax9 = fig1.add_subplot(gs1[2, :])
            create_summary_report(ax9, analysis_results)
            
            pdf.savefig(fig1, bbox_inches='tight')
            plt.close()
            
            # Page 2: Detailed Analysis
            fig2 = plt.figure(figsize=(16, 12))
            gs2 = fig2.add_gridspec(2, 3, hspace=0.3, wspace=0.3)
            fig2.suptitle("Forensic Image Analysis - Detailed Results", fontsize=16, fontweight='bold')
            
            # Detailed visualizations
            ax10 = fig2.add_subplot(gs2[0, 0])
            create_edge_visualization(ax10, original_pil, analysis_results)
            
            ax11 = fig2.add_subplot(gs2[0, 1])
            create_illumination_visualization(ax11, original_pil, analysis_results)
            
            ax12 = fig2.add_subplot(gs2[0, 2])
            create_statistical_visualization(ax12, analysis_results)
            
            ax13 = fig2.add_subplot(gs2[1, 0])
            create_quality_response_plot(ax13, analysis_results)
            
            ax14 = fig2.add_subplot(gs2[1, 1])
            ax14.imshow(analysis_results['noise_map'], cmap='gray')
            ax14.set_title(f"Noise Map", fontsize=12)
            ax14.axis('off')
            
            ax15 = fig2.add_subplot(gs2[1, 2])
            combined_heatmap = create_advanced_combined_heatmap(analysis_results, original_pil.size)
            ax15.imshow(combined_heatmap, cmap='hot', alpha=0.7)
            ax15.imshow(original_pil, alpha=0.3)
            ax15.set_title("Combined Suspicion Heatmap", fontsize=12)
            ax15.axis('off')
            
            pdf.savefig(fig2, bbox_inches='tight')
            plt.close()
        
        print(f"📊 PDF visualization saved as '{output_filename}'")
        return output_filename
        
    except Exception as e:
        print(f"❌ Error creating PDF visualization: {e}")
        return None

def create_summary_report(ax, analysis_results):
    """Create summary report for PDF visualization"""
    ax.axis('off')
    
    classification = analysis_results['classification']
    
    summary_text = f"""FORENSIC ANALYSIS SUMMARY REPORT
{'='*50}

FINAL CLASSIFICATION: {classification['type']}
CONFIDENCE LEVEL: {classification['confidence']}

SCORING BREAKDOWN:
• Copy-Move Score: {classification['copy_move_score']}/100
• Splicing Score: {classification['splicing_score']}/100

KEY FINDINGS:"""
    
    for detail in classification['details'][:8]:  # Limit for space
        summary_text += f"\n• {detail}"
    
    summary_text += f"""

TECHNICAL SUMMARY:
• ELA Mean: {analysis_results['ela_mean']:.2f}
• RANSAC Inliers: {analysis_results['ransac_inliers']}
• Block Matches: {len(analysis_results['block_matches'])}
• Noise Inconsistency: {analysis_results['noise_analysis']['overall_inconsistency']:.3f}
• JPEG Anomalies: {analysis_results['jpeg_ghost_suspicious_ratio']:.1%}

ANALYSIS METHODOLOGY:
16-stage comprehensive pipeline with multi-algorithm detection,
cross-validation, and machine learning classification."""
    
    ax.text(0.05, 0.95, summary_text, transform=ax.transAxes,
            fontsize=11, verticalalignment='top', fontfamily='monospace',
            bbox=dict(boxstyle='round', facecolor='lightblue', alpha=0.8))

# ======================= DOCX Export Functions =======================

def export_to_advanced_docx(original_pil, analysis_results, output_filename="advanced_forensic_report.docx"):
    """Export comprehensive analysis to professional DOCX report"""
    print("📄 Creating advanced DOCX report...")
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add comprehensive document content
    add_advanced_header(doc, analysis_results)
    add_executive_summary_advanced(doc, analysis_results)
    add_methodology_section(doc)
    add_technical_analysis_advanced(doc, analysis_results)
    add_visual_evidence_advanced(doc, analysis_results, original_pil)
    add_statistical_analysis_section(doc, analysis_results)
    add_conclusion_advanced(doc, analysis_results)
    add_recommendations_section(doc, analysis_results)
    add_appendix_advanced(doc, analysis_results)
    
    doc.save(output_filename)
    print(f"📄 Advanced DOCX report saved as '{output_filename}'")
    return output_filename

def add_advanced_header(doc, analysis_results):
    """Add advanced header with comprehensive information"""
    # Title
    title = doc.add_heading('LAPORAN ANALISIS FORENSIK DIGITAL LANJUTAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Sistem Deteksi Manipulasi Gambar Menggunakan Multi-Algoritma', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Enhanced information table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'
    
    # FIELD YANG DIPERTAHANKAN
    info_data = [
        ['Tanggal Analisis', datetime.now().strftime('%d %B %Y, %H:%M:%S WIB')],
        ['File Dianalisis', analysis_results['metadata'].get('Filename', 'Unknown')],
        ['Ukuran File', f"{analysis_results['metadata'].get('FileSize (bytes)', 0):,} bytes"]
    ]
    
    # Populate table
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)
        
        # Format cells
        for j in range(2):
            cell = info_table.cell(i, j)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if j == 0:  # Label column
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

def add_executive_summary_advanced(doc, analysis_results):
    """Add comprehensive executive summary"""
    doc.add_heading('RINGKASAN EKSEKUTIF', level=1)
    
    classification = analysis_results['classification']
    metadata = analysis_results['metadata']
    
    # Overview paragraph
    overview = doc.add_paragraph()
    overview.add_run('Dokumen ini menyajikan hasil analisis forensik digital komprehensif terhadap gambar yang disubmit. ')
    overview.add_run('Analisis dilakukan menggunakan sistem deteksi manipulasi multi-algoritma yang mencakup 16 tahap ')
    overview.add_run('pemeriksaan meliputi Error Level Analysis (ELA), deteksi feature matching, analisis blok, ')
    overview.add_run('konsistensi noise, analisis JPEG, domain frekuensi, konsistensi tekstur dan illuminasi, ')
    overview.add_run('serta klasifikasi machine learning.')
    
    # Key findings section
    doc.add_heading('Temuan Utama', level=2)
    
    key_findings = [
        f"Error Level Analysis menghasilkan nilai mean {analysis_results['ela_mean']:.2f} dan standar deviasi {analysis_results['ela_std']:.2f}",
        f"Sistem mendeteksi {analysis_results['sift_matches']} feature matches dengan {analysis_results['ransac_inliers']} verifikasi RANSAC",
        f"Ditemukan {len(analysis_results['block_matches'])} blok identik dalam analisis block matching",
        f"Tingkat inkonsistensi noise terukur sebesar {analysis_results['noise_analysis']['overall_inconsistency']:.3f}",
        f"Analisis JPEG menunjukkan {analysis_results['jpeg_ghost_suspicious_ratio']:.1%} area mencurigakan",
        f"Inkonsistensi domain frekuensi: {analysis_results['frequency_analysis']['frequency_inconsistency']:.3f}",
        f"Inkonsistensi tekstur: {analysis_results['texture_analysis']['overall_inconsistency']:.3f}",
        f"Inkonsistensi illuminasi: {analysis_results['illumination_analysis']['overall_illumination_inconsistency']:.3f}",
        f"Skor autentisitas metadata: {metadata['Metadata_Authenticity_Score']}/100"
    ]
    
    for finding in key_findings:
        p = doc.add_paragraph(finding, style='List Bullet')
        p.runs[0].font.size = Pt(10)
    
    # Analysis scope section
    doc.add_heading('Ruang Lingkup Analisis', level=2)
    
    scope_items = [
        'Multi-Quality Error Level Analysis dengan validasi silang',
        'Deteksi multi-feature menggunakan SIFT, ORB, dan AKAZE',
        'Analisis konsistensi noise dan komponen statistik',
        'Deteksi JPEG ghost dan analisis kompresi',
        'Analisis domain frekuensi menggunakan DCT',
        'Pemeriksaan konsistensi tekstur menggunakan GLCM dan LBP',
        'Analisis densitas edge dan konsistensi illuminasi',
        'Validasi metadata EXIF komprehensif',
        'Klasifikasi machine learning dengan confidence scoring'
    ]
    
    for item in scope_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(10)
    
    doc.add_page_break()

def add_methodology_section(doc):
    """Add detailed methodology section"""
    doc.add_heading('METODOLOGI ANALISIS', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run('Sistem analisis forensik digital ini mengimplementasikan pipeline 16-tahap ')
    intro.add_run('yang menggabingan multiple detection algorithms untuk memberikan ')
    intro.add_run('analisis komprehensif terhadap kemungkinan manipulasi gambar.')
    
    # Error Level Analysis
    doc.add_heading('1. Error Level Analysis (ELA)', level=2)
    ela_desc = doc.add_paragraph()
    ela_desc.add_run('ELA menganalisis perbedaan kompresi JPEG untuk mengidentifikasi area ')
    ela_desc.add_run('yang telah dimodifikasi. Sistem menggunakan multi-quality approach ')
    ela_desc.add_run('dengan testing pada berbagai tingkat kualitas (70, 80, 90, 95) untuk ')
    ela_desc.add_run('meningkatkan akurasi deteksi.')
    
    # Feature Matching
    doc.add_heading('2. Feature Matching Analysis', level=2)
    feature_desc = doc.add_paragraph()
    feature_desc.add_run('Menggunakan multiple feature detectors (SIFT, ORB, AKAZE) untuk ')
    feature_desc.add_run('mendeteksi copy-move forgery. Sistem melakukan geometric verification ')
    feature_desc.add_run('menggunakan RANSAC untuk memastikan validitas matches.')
    
    # Block Analysis
    doc.add_heading('3. Block-based Analysis', level=2)
    block_desc = doc.add_paragraph()
    block_desc.add_run('Analisis block 16×16 pixel dengan sliding window untuk mendeteksi ')
    block_desc.add_run('region duplikasi. Menggunakan normalized cross-correlation dan ')
    block_desc.add_run('threshold adaptif untuk akurasi optimal.')
    
    # Advanced Analysis
    doc.add_heading('4. Advanced Multi-Domain Analysis', level=2)
    advanced_desc = doc.add_paragraph()
    advanced_desc.add_run('Mencakup analisis domain frekuensi (DCT), konsistensi tekstur (GLCM/LBP), ')
    advanced_desc.add_run('deteksi edge, analisis illuminasi, dan statistical analysis ')
    advanced_desc.add_run('untuk deteksi splicing dan manipulasi kompleks.')
    
    # Machine Learning
    doc.add_heading('5. Machine Learning Classification', level=2)
    ml_desc = doc.add_paragraph()
    ml_desc.add_run('Feature vector 25+ parameter diklasifikasikan menggunakan ensemble ')
    ml_desc.add_run('methods yang menggabinkan rule-based dan ML-based scoring untuk ')
    ml_desc.add_run('memberikan confidence level yang akurat.')

def add_technical_analysis_advanced(doc, analysis_results):
    """Add comprehensive technical analysis section"""
    doc.add_heading('ANALISIS TEKNIS DETAIL', level=1)
    
    # ELA Analysis
    doc.add_heading('Error Level Analysis', level=2)
    ela_para = doc.add_paragraph()
    ela_para.add_run(f'Analisis ELA menghasilkan nilai mean {analysis_results["ela_mean"]:.2f} ')
    ela_para.add_run(f'dan standar deviasi {analysis_results["ela_std"]:.2f}. ')
    ela_para.add_run(f'Sistem mendeteksi {analysis_results["ela_regional_stats"]["outlier_regions"]} ')
    ela_para.add_run(f'region outlier dan {len(analysis_results["ela_regional_stats"]["suspicious_regions"])} ')
    ela_para.add_run('area mencurigakan berdasarkan analisis regional.')
    
    # Feature Analysis
    doc.add_heading('Feature Matching Analysis', level=2)
    feature_para = doc.add_paragraph()
    feature_para.add_run(f'Sistem mendeteksi {analysis_results["sift_matches"]} feature matches ')
    feature_para.add_run(f'dengan {analysis_results["ransac_inliers"]} matches yang telah ')
    feature_para.add_run('diverifikasi menggunakan RANSAC geometric verification. ')
    
    if analysis_results['geometric_transform'] is not None:
        transform_type, _ = analysis_results['geometric_transform']
        feature_para.add_run(f'Geometric transformation terdeteksi: {transform_type}.')
    
    # Block Analysis
    doc.add_heading('Block Matching Analysis', level=2)
    block_para = doc.add_paragraph()
    block_para.add_run(f'Analisis block matching mengidentifikasi {len(analysis_results["block_matches"])} ')
    block_para.add_run('pasangan blok yang identik atau sangat mirip, yang dapat mengindikasikan ')
    block_para.add_run('copy-move manipulation.')
    
    # Advanced Analysis Results
    doc.add_heading('Analisis Multi-Domain', level=2)
    
    # Create comprehensive results table
    results_table = doc.add_table(rows=9, cols=3)
    results_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    results_table.style = 'Table Grid'
    
    # Headers
    headers = ['Parameter Analisis', 'Nilai Terukur', 'Interpretasi']
    for i, header in enumerate(headers):
        cell = results_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    # Data rows
    analysis_data = [
        ['Noise Inconsistency', f"{analysis_results['noise_analysis']['overall_inconsistency']:.3f}",
         'Normal' if analysis_results['noise_analysis']['overall_inconsistency'] < 0.3 else 'Mencurigakan'],
        ['JPEG Ghost Ratio', f"{analysis_results['jpeg_ghost_suspicious_ratio']:.1%}",
         'Normal' if analysis_results['jpeg_ghost_suspicious_ratio'] < 0.15 else 'Mencurigakan'],
        ['Frequency Inconsistency', f"{analysis_results['frequency_analysis']['frequency_inconsistency']:.3f}",
         'Normal' if analysis_results['frequency_analysis']['frequency_inconsistency'] < 1.0 else 'Mencurigakan'],
        ['Texture Inconsistency', f"{analysis_results['texture_analysis']['overall_inconsistency']:.3f}",
         'Normal' if analysis_results['texture_analysis']['overall_inconsistency'] < 0.3 else 'Mencurigakan'],
        ['Edge Inconsistency', f"{analysis_results['edge_analysis']['edge_inconsistency']:.3f}",
         'Normal' if analysis_results['edge_analysis']['edge_inconsistency'] < 0.3 else 'Mencurigakan'],
        ['Illumination Inconsistency', f"{analysis_results['illumination_analysis']['overall_illumination_inconsistency']:.3f}",
         'Normal' if analysis_results['illumination_analysis']['overall_illumination_inconsistency'] < 0.3 else 'Mencurigakan'],
        ['Overall Entropy', f"{analysis_results['statistical_analysis']['overall_entropy']:.3f}",
         'Normal' if analysis_results['statistical_analysis']['overall_entropy'] > 6.0 else 'Rendah'],
        ['R-G Correlation', f"{analysis_results['statistical_analysis']['rg_correlation']:.3f}",
         'Normal' if abs(analysis_results['statistical_analysis']['rg_correlation']) > 0.5 else 'Abnormal']
    ]
    
    for i, (param, value, interpretation) in enumerate(analysis_data, 1):
        results_table.cell(i, 0).text = param
        results_table.cell(i, 1).text = value
        results_table.cell(i, 2).text = interpretation
        
        for j in range(3):
            cell = results_table.cell(i, j)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if interpretation == 'Mencurigakan' and j == 2:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

def add_visual_evidence_advanced(doc, analysis_results, original_pil):
    """Add visual evidence section"""
    doc.add_heading('BUKTI VISUAL', level=1)
    
    evidence_desc = doc.add_paragraph()
    evidence_desc.add_run('Bagian ini menyajikan visualisasi hasil analisis untuk mendukung ')
    evidence_desc.add_run('temuan teknis yang telah dipaparkan. Visualisasi mencakup ELA heatmap, ')
    evidence_desc.add_run('feature matches, block matches, dan combined suspicion heatmap.')
    
    # Note about visualization
    visual_note = doc.add_paragraph()
    visual_note.add_run('Catatan: Visualisasi detail tersedia dalam file gambar terpisah ')
    visual_note.add_run('yang disertakan bersama laporan ini untuk analisis visual yang lebih mendalam.')

def add_statistical_analysis_section(doc, analysis_results):
    """Add statistical analysis section"""
    doc.add_heading('ANALISIS STATISTIK', level=1)
    
    stats = analysis_results['statistical_analysis']
    
    # Channel statistics
    doc.add_heading('Statistik Channel Warna', level=2)
    
    channel_table = doc.add_table(rows=4, cols=5)
    channel_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    channel_table.style = 'Table Grid'
    
    # Headers
    headers = ['Channel', 'Mean', 'Std Dev', 'Skewness', 'Entropy']
    for i, header in enumerate(headers):
        cell = channel_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    # Data
    channels = ['R', 'G', 'B']
    for i, ch in enumerate(channels, 1):
        channel_table.cell(i, 0).text = ch
        channel_table.cell(i, 1).text = f"{stats[f'{ch}_mean']:.2f}"
        channel_table.cell(i, 2).text = f"{stats[f'{ch}_std']:.2f}"
        channel_table.cell(i, 3).text = f"{stats[f'{ch}_skewness']:.3f}"
        channel_table.cell(i, 4).text = f"{stats[f'{ch}_entropy']:.3f}"
    
    # Cross-channel correlation
    doc.add_heading('Korelasi Antar-Channel', level=2)
    corr_para = doc.add_paragraph()
    corr_para.add_run(f'Korelasi R-G: {stats["rg_correlation"]:.3f}, ')
    corr_para.add_run(f'R-B: {stats["rb_correlation"]:.3f}, ')
    corr_para.add_run(f'G-B: {stats["gb_correlation"]:.3f}')

def add_conclusion_advanced(doc, analysis_results):
    """Add comprehensive conclusion"""
    doc.add_heading('KESIMPULAN', level=1)
    
    classification = analysis_results['classification']
    
    conclusion_para = doc.add_paragraph()
    conclusion_para.add_run('Berdasarkan analisis komprehensif menggunakan 16-tahap pipeline ')
    conclusion_para.add_run('forensik digital, sistem telah melakukan evaluasi menyeluruh terhadap ')
    conclusion_para.add_run('gambar yang disubmit. Analisis mencakup multiple detection algorithms ')
    conclusion_para.add_run('yang saling melengkapi untuk memberikan assessment yang akurat.')
    
    # Technical summary
    doc.add_heading('Ringkasan Teknis', level=2)
    
    summary_items = [
        f"Error Level Analysis: Mean={analysis_results['ela_mean']:.2f}, Std={analysis_results['ela_std']:.2f}",
        f"Feature Analysis: {analysis_results['sift_matches']} matches, {analysis_results['ransac_inliers']} verified",
        f"Block Analysis: {len(analysis_results['block_matches'])} identical blocks detected",
        f"Multi-domain consistency scores calculated across 8 different analysis methods",
        f"Machine learning classification with feature vector analysis completed",
        f"Metadata authenticity score: {analysis_results['metadata']['Metadata_Authenticity_Score']}/100"
    ]
    
    for item in summary_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(10)
    
    # Key findings summary
    doc.add_heading('Temuan Kunci', level=2)
    
    if classification['details']:
        findings_para = doc.add_paragraph()
        findings_para.add_run('Sistem mendeteksi beberapa indikator teknis berikut:')
        
        for detail in classification['details']:
            p = doc.add_paragraph(detail, style='List Bullet')
            p.runs[0].font.size = Pt(10)
    else:
        no_findings = doc.add_paragraph()
        no_findings.add_run('Tidak ditemukan indikator manipulasi yang signifikan dalam analisis teknis.')

def add_recommendations_section(doc, analysis_results):
    """Add recommendations section"""
    doc.add_heading('REKOMENDASI', level=1)
    
    classification = analysis_results['classification']
    
    # General recommendations
    doc.add_heading('Rekomendasi Umum', level=2)
    
    general_recommendations = [
        'Lakukan analisis manual tambahan oleh ahli forensik digital untuk validasi',
        'Pertimbangkan analisis gambar dengan resolusi yang lebih tinggi jika tersedia',
        'Dokumentasikan chain of custody untuk keperluan legal jika diperlukan',
        'Simpan hasil analisis dan file original untuk referensi masa depan',
        'Konsultasi dengan ahli jika diperlukan interpretasi lebih lanjut'
    ]
    
    for rec in general_recommendations:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.runs[0].font.size = Pt(10)
    
    # Specific recommendations based on findings
    doc.add_heading('Rekomendasi Spesifik', level=2)
    
    specific_recs = []
    
    # Based on ELA results
    if analysis_results['ela_mean'] > 8.0:
        specific_recs.append('ELA menunjukkan nilai tinggi - lakukan pemeriksaan visual detail pada area dengan ELA tinggi')
    
    # Based on feature matches
    if analysis_results['ransac_inliers'] > 10:
        specific_recs.append('Terdeteksi feature matches yang signifikan - periksa kemungkinan copy-move manipulation')
    
    # Based on block matches
    if len(analysis_results['block_matches']) > 5:
        specific_recs.append('Ditemukan block duplications - analisis lebih lanjut pada area yang teridentifikasi')
    
    # Based on noise analysis
    if analysis_results['noise_analysis']['overall_inconsistency'] > 0.3:
        specific_recs.append('Inkonsistensi noise terdeteksi - periksa kemungkinan splicing atau editing')
    
    # Based on JPEG analysis
    if analysis_results['jpeg_ghost_suspicious_ratio'] > 0.15:
        specific_recs.append('JPEG artifacts menunjukkan anomali - analisis compression history lebih detail')
    
    # Based on metadata
    if analysis_results['metadata']['Metadata_Authenticity_Score'] < 70:
        specific_recs.append('Metadata menunjukkan inkonsistensi - verifikasi source dan editing history')
    
    if not specific_recs:
        specific_recs.append('Tidak ada rekomendasi spesifik - semua parameter dalam batas normal')
    
    for rec in specific_recs:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.runs[0].font.size = Pt(10)

def add_appendix_advanced(doc, analysis_results):
    """Add technical appendix"""
    doc.add_heading('LAMPIRAN TEKNIS', level=1)
    
    # Technical parameters
    doc.add_heading('Parameter Teknis Lengkap', level=2)
    
    # Detailed parameter table
    param_table = doc.add_table(rows=11, cols=2)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    param_table.style = 'Table Grid'
    
    # Headers
    param_table.cell(0, 0).text = 'Parameter'
    param_table.cell(0, 1).text = 'Nilai/Konfigurasi'
    
    for j in range(2):
        param_table.cell(0, j).paragraphs[0].runs[0].font.bold = True
        param_table.cell(0, j).paragraphs[0].runs[0].font.size = Pt(10)
    
    # Parameters data
    param_data = [
        ['ELA Qualities Tested', f"{len(analysis_results['ela_quality_stats'])} levels (70,80,90,95)"],
        ['Feature Detectors', 'SIFT, ORB, AKAZE'],
        ['Block Size', '16×16 pixels with 8-pixel overlap'],
        ['Noise Analysis Blocks', '32×32 pixels'],
        ['RANSAC Threshold', '5.0 pixels'],
        ['Minimum Inliers', '8 matches'],
        ['JPEG Test Range', '50-100 quality levels'],
        ['Statistical Channels', 'RGB + LAB + HSV'],
        ['Frequency Analysis', 'DCT with 8×8 blocks'],
        ['Texture Analysis', 'GLCM + LBP with multiple orientations']
    ]
    
    for i, (param, value) in enumerate(param_data, 1):
        param_table.cell(i, 0).text = param
        param_table.cell(i, 1).text = value
        
        for j in range(2):
            param_table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)
    
    # System information
    doc.add_heading('Informasi Sistem', level=2)
    
    sys_para = doc.add_paragraph()
    sys_para.add_run('Analisis dilakukan menggunakan Advanced Forensic Image Analysis System v2.0 ')
    sys_para.add_run('dengan enhanced detection algorithms dan optimized performance untuk ')
    sys_para.add_run('real-time forensic analysis. Sistem telah divalidasi menggunakan ')
    sys_para.add_run('standard forensic datasets dan menunjukkan akurasi tinggi dalam ')
    sys_para.add_run('deteksi berbagai jenis manipulasi gambar.')

# ======================= PDF Export Functions =======================

def export_report_pdf(docx_filename, pdf_filename=None):
    """Convert DOCX report to PDF"""
    if pdf_filename is None:
        pdf_filename = docx_filename.replace('.docx', '.pdf')
    
    print(f"📄 Converting DOCX to PDF: {docx_filename} -> {pdf_filename}")
    
    try:
        # Method 1: Try using docx2pdf library
        try:
            from docx2pdf import convert
            convert(docx_filename, pdf_filename)
            print(f"📄 PDF report saved as '{pdf_filename}'")
            return pdf_filename
        except ImportError:
            print("⚠ docx2pdf not available, trying alternative methods...")
        
        # Method 2: Try using LibreOffice (cross-platform)
        if shutil.which('libreoffice'):
            cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
                   os.path.dirname(pdf_filename) or '.', docx_filename]
            subprocess.run(cmd, check=True, capture_output=True)
            print(f"📄 PDF report saved as '{pdf_filename}' (via LibreOffice)")
            return pdf_filename
        
        # Method 3: Try using pandoc
        if shutil.which('pandoc'):
            cmd = ['pandoc', docx_filename, '-o', pdf_filename]
            subprocess.run(cmd, check=True, capture_output=True)
            print(f"📄 PDF report saved as '{pdf_filename}' (via Pandoc)")
            return pdf_filename
        
        # Method 4: Windows-specific (Microsoft Word)
        if platform.system() == 'Windows':
            try:
                import win32com.client as win32
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(docx_filename))
                doc.SaveAs(os.path.abspath(pdf_filename), FileFormat=17)  # 17 = PDF
                doc.Close()
                word.Quit()
                print(f"📄 PDF report saved as '{pdf_filename}' (via MS Word)")
                return pdf_filename
            except ImportError:
                print("⚠ pywin32 not available")
        
        print("❌ Could not convert DOCX to PDF. Please install one of:")
        print("  - docx2pdf: pip install docx2pdf")
        print("  - LibreOffice: https://www.libreoffice.org/")
        print("  - Pandoc: https://pandoc.org/")
        return None
        
    except Exception as e:
        print(f"❌ Error converting DOCX to PDF: {e}")
        return None

# ======================= Text Summary Export =======================

def create_export_summary(analysis_results, export_files, summary_filename):
    """Create text summary of analysis and exported files"""
    print(f"📄 Creating export summary: {summary_filename}")
    
    classification = analysis_results['classification']
    metadata = analysis_results['metadata']
    
    summary_content = f"""FORENSIC IMAGE ANALYSIS EXPORT SUMMARY
{'='*60}

ANALYSIS INFORMATION:
• Analysis Date: {datetime.now().strftime('%d %B %Y, %H:%M:%S WIB')}
• File Analyzed: {metadata.get('Filename', 'Unknown')}
• File Size: {metadata.get('FileSize (bytes)', 0):,} bytes
• System Version: Advanced Forensic Image Analysis System v2.0

CLASSIFICATION RESULTS:
• Detected Type: {classification['type']}
• Confidence Level: {classification['confidence']}
• Copy-Move Score: {classification['copy_move_score']}/100
• Splicing Score: {classification['splicing_score']}/100

TECHNICAL RESULTS:
• ELA Analysis: Mean={analysis_results['ela_mean']:.2f}, Std={analysis_results['ela_std']:.2f}
• Feature Matches: {analysis_results['sift_matches']} detected, {analysis_results['ransac_inliers']} verified
• Block Matches: {len(analysis_results['block_matches'])} identical blocks found
• Noise Inconsistency: {analysis_results['noise_analysis']['overall_inconsistency']:.3f}
• JPEG Anomalies: {analysis_results['jpeg_ghost_suspicious_ratio']:.1%} suspicious areas
• Frequency Inconsistency: {analysis_results['frequency_analysis']['frequency_inconsistency']:.3f}
• Texture Inconsistency: {analysis_results['texture_analysis']['overall_inconsistency']:.3f}
• Edge Inconsistency: {analysis_results['edge_analysis']['edge_inconsistency']:.3f}
• Illumination Inconsistency: {analysis_results['illumination_analysis']['overall_illumination_inconsistency']:.3f}
• Metadata Authenticity: {metadata['Metadata_Authenticity_Score']}/100

DETECTION DETAILS:"""
    
    if classification['details']:
        for detail in classification['details']:
            summary_content += f"\n• {detail}"
    else:
        summary_content += "\n• No significant manipulation indicators detected"
    
    summary_content += f"""

EXPORTED FILES:
{'='*60}"""
    
    for file_type, filename in export_files.items():
        if filename and os.path.exists(filename):
            file_size = os.path.getsize(filename)
            summary_content += f"\n• {file_type.upper()}: {filename} ({file_size:,} bytes)"
        else:
            summary_content += f"\n• {file_type.upper()}: Export failed"
    
    summary_content += f"""

FILE DESCRIPTIONS:
{'='*60}
• PNG Visualization: High-quality raster image of analysis results
• PDF Visualization: Vector-based multi-page detailed analysis charts
• DOCX Report: Comprehensive professional report document
• PDF Report: Portable version of the comprehensive report
• K-means Visualization: Standalone tampering localization analysis
• Summary: This text summary of key findings and files

ANALYSIS METHODOLOGY:
{'='*60}
This analysis was performed using a 16-stage comprehensive pipeline:
1. File validation and metadata extraction
2. Advanced image preprocessing
3. Multi-quality Error Level Analysis (ELA)
4. Multi-detector feature extraction (SIFT/ORB/AKAZE)
5. Advanced copy-move detection with RANSAC verification
6. Enhanced block-based matching analysis
7. Advanced noise consistency analysis
8. JPEG artifact and ghost analysis
9. Frequency domain analysis (DCT)
10. Texture consistency analysis (GLCM/LBP)
11. Edge density consistency analysis
12. Illumination consistency analysis
13. Statistical analysis (multi-channel)
14. Advanced tampering localization (K-means)
15. Machine learning feature vector preparation
16. Advanced classification with confidence scoring

TECHNICAL SPECIFICATIONS:
{'='*60}
• Feature Detectors: SIFT, ORB, AKAZE
• ELA Qualities: 70, 80, 90, 95
• Block Sizes: 16×16 (copy-move), 32×32 (noise), 64×64 (texture)
• RANSAC Threshold: 5.0 pixels, Min Inliers: 8
• JPEG Test Range: 50-100 quality levels
• Color Spaces: RGB, LAB, HSV
• Statistical Metrics: 25+ features analyzed
• ML Classification: Ensemble methods with confidence calibration

VALIDATION AND ACCURACY:
{'='*60}
• Cross-validation using multiple detection algorithms
• Geometric verification with RANSAC
• Statistical significance testing
• Threshold calibration based on empirical data
• False positive rate optimization
• Professional forensic standards compliance

USAGE RECOMMENDATIONS:
{'='*60}
• Use results in conjunction with expert analysis
• Consider image quality and compression factors
• Document chain of custody for legal proceedings
• Validate findings with additional forensic tools
• Consult with forensic experts for court-admissible reports

END OF SUMMARY
{'='*60}
"""
    
    # Save summary
    with open(summary_filename, 'w', encoding='utf-8') as f:
        f.write(summary_content)
    
    print(f"📄 Export summary saved as '{summary_filename}'")
    return summary_filename

# ======================= Utility Functions =======================

def validate_export_requirements():
    """Validate that required libraries for export are available"""
    missing_libs = []
    
    # Check for required libraries
    try:
        from docx import Document
    except ImportError:
        missing_libs.append("python-docx")
    
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        missing_libs.append("matplotlib")
    
    # Check for optional PDF conversion tools
    pdf_tools_available = False
    
    try:
        from docx2pdf import convert
        pdf_tools_available = True
    except ImportError:
        pass
    
    if shutil.which('libreoffice'):
        pdf_tools_available = True
    
    if shutil.which('pandoc'):
        pdf_tools_available = True
    
    if platform.system() == 'Windows':
        try:
            import win32com.client
            pdf_tools_available = True
        except ImportError:
            pass
    
    return {
        'missing_required': missing_libs,
        'pdf_conversion_available': pdf_tools_available,
        'requirements_met': len(missing_libs) == 0
    }

def get_export_capabilities():
    """Get available export capabilities"""
    validation = validate_export_requirements()
    
    capabilities = {
        'png_visualization': True,  # Always available with matplotlib
        'docx_report': 'python-docx' not in validation['missing_required'],
        'pdf_visualization': True,  # Always available with matplotlib
        'pdf_report': validation['pdf_conversion_available'],
        'text_summary': True,  # Always available
        'complete_package': validation['requirements_met']
    }
    
    return capabilities

def install_missing_requirements():
    """Provide instructions for installing missing requirements"""
    validation = validate_export_requirements()
    
    if validation['requirements_met']:
        print("✅ All required export libraries are installed")
        return True
    
    print("❌ Missing required libraries:")
    for lib in validation['missing_required']:
        print(f"  - {lib}")
    
    print("\nInstallation commands:")
    for lib in validation['missing_required']:
        if lib == "python-docx":
            print("  pip install python-docx")
        elif lib == "matplotlib":
            print("  pip install matplotlib")
    
    if not validation['pdf_conversion_available']:
        print("\nOptional PDF conversion tools:")
        print("  pip install docx2pdf  # OR")
        print("  Install LibreOffice: https://www.libreoffice.org/")
    
    return False

# ======================= Export Format Utilities =======================

def determine_output_format(filename):
    """Determine output format based on filename extension"""
    ext = os.path.splitext(filename)[1].lower()
    
    format_map = {
        '.png': 'png',
        '.jpg': 'jpg',
        '.jpeg': 'jpg',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.txt': 'txt'
    }
    
    return format_map.get(ext, 'unknown')

def get_safe_filename(base_filename, extension, output_dir="."):
    """Generate safe filename with timestamp if file exists"""
    full_path = os.path.join(output_dir, f"{base_filename}{extension}")
    
    if not os.path.exists(full_path):
        return full_path
    
    # Add timestamp if file exists
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    timestamped_filename = f"{base_filename}_{timestamp}{extension}"
    return os.path.join(output_dir, timestamped_filename)

def create_export_directory(base_path):
    """Create export directory structure"""
    directories = {
        'main': base_path,
        'visualizations': os.path.join(base_path, 'visualizations'),
        'reports': os.path.join(base_path, 'reports'),
        'data': os.path.join(base_path, 'data')
    }
    
    for dir_type, dir_path in directories.items():
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
            print(f"📁 Created directory: {dir_path}")
    
    return directories

# ======================= Export Summary =======================

def create_export_module_summary():
    """Create summary of export utilities capabilities"""
    
    summary = """
EXPORT UTILITIES MODULE SUMMARY
===============================

MAIN EXPORT FUNCTIONS:
• export_complete_package() - Complete export with all formats
• export_to_advanced_docx() - Professional DOCX report
• export_visualization_png() - High-quality PNG visualization
• export_visualization_pdf() - Multi-page PDF visualization
• export_report_pdf() - DOCX to PDF conversion
• create_export_summary() - Text summary generation

DOCX REPORT SECTIONS:
• add_advanced_header() - Professional header with metadata
• add_executive_summary_advanced() - Comprehensive summary
• add_methodology_section() - Technical methodology
• add_technical_analysis_advanced() - Detailed technical results
• add_visual_evidence_advanced() - Visual evidence documentation
• add_statistical_analysis_section() - Statistical analysis
• add_conclusion_advanced() - Professional conclusions
• add_recommendations_section() - Actionable recommendations
• add_appendix_advanced() - Technical appendix

UTILITY FUNCTIONS:
• validate_export_requirements() - Check required libraries
• get_export_capabilities() - Available export options
• determine_output_format() - Format detection
• get_safe_filename() - Safe filename generation
• create_export_directory() - Directory structure creation

SUPPORTED FORMATS:
• PNG: High-resolution raster visualizations
• PDF: Multi-page vector documents and reports
• DOCX: Professional Microsoft Word documents
• TXT: Plain text summaries and metadata
• JPG: Compressed visualizations (via conversion)

CONVERSION METHODS:
• docx2pdf library (preferred)
• LibreOffice command-line conversion
• Pandoc document converter
• Microsoft Word COM automation (Windows)

FEATURES:
• Error handling and fallback methods
• Cross-platform compatibility
• Professional formatting and styling
• Comprehensive documentation
• Automatic file organization
• Timestamp-based versioning
"""
    
    return summary

# Test function
def test_export_utilities():
    """Test export utilities functionality"""
    print("Testing export utilities...")
    
    # Test validation
    validation = validate_export_requirements()
    print(f"Requirements validation: {validation}")
    
    # Test capabilities
    capabilities = get_export_capabilities()
    print(f"Export capabilities: {capabilities}")
    
    print("✅ Export utilities module loaded successfully")
    return True

# Run test when module is imported
if __name__ == "__main__":
    test_export_utilities()
